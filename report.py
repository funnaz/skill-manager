"""Export scan reports in multiple languages and formats."""

from __future__ import annotations

import io
import json
from datetime import datetime
from pathlib import Path
from typing import Any

from constants import GITHUB_URL
from scanner import scan_all

EXPORT_FORMATS = ("md", "docx", "pdf")
EXPORT_LANGS = ("zh", "en")
FMT_ALIASES = {
    "word": "docx",
    "markdown": "md",
    ".docx": "docx",
    ".md": "md",
    ".pdf": "pdf",
}


def normalize_export_fmt(fmt: str) -> str:
    return FMT_ALIASES.get(fmt.lower().strip(), fmt.lower().strip())


def _category_counts(data: dict[str, Any]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for skill in data["skills"]:
        counts[skill["category"]] = counts.get(skill["category"], 0) + 1
    return counts


def _deletable_counts(data: dict[str, Any]) -> tuple[int, int]:
    deletable = sum(1 for skill in data["skills"] if skill.get("deletable"))
    return deletable, len(data["skills"]) - deletable


def _top_agents(data: dict[str, Any]) -> list[tuple[str, int]]:
    return sorted(
        ((agent["label"], agent["installed"]) for agent in data["agents"] if agent["configured"]),
        key=lambda item: item[1],
        reverse=True,
    )


def _build_analysis(lang: str, data: dict[str, Any]) -> list[dict[str, Any]]:
    totals = data["totals"]
    categories = _category_counts(data)
    deletable, protected = _deletable_counts(data)
    agents = _top_agents(data)
    shared_ratio = round(totals["shared_skills"] / totals["skills"] * 100) if totals["skills"] else 0

    if lang == "zh":
        return [
            {
                "title": "概览解读",
                "paragraphs": [
                    f"本机共扫描到 **{totals['skills']}** 个 Skill，其中 **{totals['shared_skills']}** 个为多 Agent 共享（约占 {shared_ratio}%）。"
                    f"当前已配置 **{totals['agents_configured']}** 个 Agent 环境，Grok 侧禁用 **{totals['disabled']}** 个。",
                    "共享目录 `.agents/skills` 是跨 Agent 复用的核心资产区；Grok 用户目录与内置目录则体现运行时自带能力。",
                ],
            },
            {
                "title": "Agent 生态",
                "paragraphs": [
                    "、".join(f"{label} {count} 个" for label, count in agents) or "暂无已配置的 Agent。",
                    "若同一 Skill 同时出现在 Grok 与共享目录，通常意味着你在主动做跨平台能力对齐。",
                ],
            },
            {
                "title": "分类结构",
                "paragraphs": [
                    "、".join(f"{cat} {count} 个" for cat, count in sorted(categories.items(), key=lambda x: -x[1])),
                    f"可删除 {deletable} 个，受保护 {protected} 个。内置、Marketplace 与包管理类 Skill 不建议手动删除。",
                ],
            },
            {
                "title": "维护建议",
                "paragraphs": [
                    "定期执行「检查更新」，优先处理「安装版本落后」项；对你改过功能的 Skill 使用「整合更新」而非覆盖升级。",
                    "对长期未使用的用户自定义 Skill，可在面板中批量勾选后删除，减少扫描噪音与路径冲突。",
                    "建议将高价值自定义 Skill 同步到 GitHub 私有仓库，便于换机恢复与版本追踪。",
                ],
            },
        ]

    return [
        {
            "title": "Overview",
            "paragraphs": [
                f"This machine has **{totals['skills']}** skills, including **{totals['shared_skills']}** shared across agents (~{shared_ratio}%).",
                f"**{totals['agents_configured']}** agent environments are configured; **{totals['disabled']}** skills are disabled in Grok.",
                "The shared `.agents/skills` directory is the main cross-agent asset pool; Grok user and bundled dirs reflect runtime capabilities.",
            ],
        },
        {
            "title": "Agent Ecosystem",
            "paragraphs": [
                ", ".join(f"{label}: {count}" for label, count in agents) or "No configured agents detected.",
                "When a skill appears in both Grok and the shared directory, it usually indicates intentional cross-platform alignment.",
            ],
        },
        {
            "title": "Category Breakdown",
            "paragraphs": [
                ", ".join(f"{cat}: {count}" for cat, count in sorted(categories.items(), key=lambda x: -x[1])),
                f"{deletable} skills are deletable; {protected} are protected. Bundled, marketplace, and package-managed skills should not be removed manually.",
            ],
        },
        {
            "title": "Maintenance Recommendations",
            "paragraphs": [
                "Run update checks regularly; upgrade outdated installs first. Use merge updates for skills you have customized.",
                "Batch-delete unused custom skills to reduce scan noise and path conflicts.",
                "Mirror high-value custom skills to a private GitHub repo for recovery and version tracking.",
            ],
        },
    ]


def _report_meta(lang: str, data: dict[str, Any]) -> dict[str, str]:
    if lang == "zh":
        return {
            "title": "Skill 管家扫描报告",
            "generated": "生成时间",
            "home": "用户目录",
            "project": "项目地址",
            "summary": "数据摘要",
            "analysis": "分析解读",
            "agents": "Agent 概览",
            "skills": "Skill 清单",
            "headers_agents": ["Agent", "已安装", "已配置", "根目录"],
            "headers_skills": ["名称", "分类", "Agent", "可删除", "路径"],
            "yes": "是",
            "no": "否",
            "total_skills": "Skill 总数",
            "configured_agents": "已配置 Agent",
            "shared_skills": "共享 Skill",
            "disabled": "已禁用",
        }
    return {
        "title": "Skill Manager Scan Report",
        "generated": "Generated",
        "home": "Home",
        "project": "Project",
        "summary": "Summary",
        "analysis": "Analysis",
        "agents": "Agents",
        "skills": "Skill Inventory",
        "headers_agents": ["Agent", "Installed", "Configured", "Roots"],
        "headers_skills": ["Name", "Category", "Agents", "Deletable", "Path"],
        "yes": "yes",
        "no": "no",
        "total_skills": "Total Skills",
        "configured_agents": "Configured Agents",
        "shared_skills": "Shared Skills",
        "disabled": "Disabled",
    }


def _markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    head = "| " + " | ".join(headers) + " |"
    sep = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = "\n".join("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join([head, sep, body])


def build_markdown_report(data: dict[str, Any] | None = None, lang: str = "en") -> str:
    data = data or scan_all()
    meta = _report_meta(lang, data)
    analysis = _build_analysis(lang, data)
    lines = [
        f"# {meta['title']}",
        "",
        f"- {meta['generated']}: {data['scanned_at']}",
        f"- {meta['home']}: `{data['home']}`",
        f"- {meta['project']}: {GITHUB_URL}",
        "",
        f"## {meta['summary']}",
        "",
        f"- {meta['total_skills']}: **{data['totals']['skills']}**",
        f"- {meta['configured_agents']}: **{data['totals']['agents_configured']}**",
        f"- {meta['shared_skills']}: **{data['totals']['shared_skills']}**",
        f"- {meta['disabled']}: **{data['totals']['disabled']}**",
        "",
        f"## {meta['analysis']}",
        "",
    ]
    for section in analysis:
        lines.append(f"### {section['title']}")
        lines.append("")
        lines.extend(section["paragraphs"])
        lines.append("")

    lines.extend([
        f"## {meta['agents']}",
        "",
        _markdown_table(
            meta["headers_agents"],
            [
                [
                    agent["label"],
                    str(agent["installed"]),
                    meta["yes"] if agent["configured"] else meta["no"],
                    "<br>".join(agent["roots"]) or "-",
                ]
                for agent in data["agents"]
            ],
        ),
        "",
        f"## {meta['skills']}",
        "",
        _markdown_table(
            meta["headers_skills"],
            [
                [
                    skill["name"],
                    skill["category"],
                    ", ".join(skill["agent_labels"]) or "-",
                    meta["yes"] if skill["deletable"] else meta["no"],
                    skill["resolved_path"],
                ]
                for skill in data["skills"]
            ],
        ),
        "",
    ])
    return "\n".join(lines)


def _pdf_font_path() -> Path | None:
    candidates = [
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simfang.ttf"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return path
    return None


def build_pdf_report(data: dict[str, Any] | None = None, lang: str = "zh") -> bytes:
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise ValueError("PDF 导出需要安装 fpdf2：pip install fpdf2") from exc

    data = data or scan_all()
    meta = _report_meta(lang, data)
    analysis = _build_analysis(lang, data)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    font_path = _pdf_font_path()
    if font_path:
        pdf.add_font("ReportFont", "", str(font_path))
        font_family = "ReportFont"
    elif lang == "zh":
        raise ValueError("未找到中文字体（simhei/msyh），无法导出中文 PDF")
    else:
        font_family = "Helvetica"

    def write_line(text: str, size: int = 11, gap: int = 6) -> None:
        pdf.set_font(font_family, size=size)
        pdf.multi_cell(0, gap, text)
        pdf.ln(2)

    write_line(meta["title"], size=16, gap=8)
    write_line(f"{meta['generated']}: {data['scanned_at']}")
    write_line(f"{meta['home']}: {data['home']}")
    write_line(f"{meta['project']}: {GITHUB_URL}")
    pdf.ln(4)

    write_line(meta["summary"], size=13, gap=7)
    write_line(f"{meta['total_skills']}: {data['totals']['skills']}")
    write_line(f"{meta['configured_agents']}: {data['totals']['agents_configured']}")
    write_line(f"{meta['shared_skills']}: {data['totals']['shared_skills']}")
    write_line(f"{meta['disabled']}: {data['totals']['disabled']}")
    pdf.ln(4)

    write_line(meta["analysis"], size=13, gap=7)
    for section in analysis:
        write_line(section["title"], size=12, gap=6)
        for paragraph in section["paragraphs"]:
            clean = paragraph.replace("**", "")
            write_line(clean)
        pdf.ln(2)

    write_line(meta["agents"], size=13, gap=7)
    for agent in data["agents"]:
        roots = "; ".join(agent["roots"]) or "-"
        write_line(
            f"{agent['label']} | {agent['installed']} | "
            f"{meta['yes'] if agent['configured'] else meta['no']} | {roots}",
            size=10,
            gap=5,
        )
    pdf.ln(4)

    write_line(meta["skills"], size=13, gap=7)
    for skill in data["skills"]:
        agents = ", ".join(skill["agent_labels"]) or "-"
        write_line(
            f"{skill['name']} | {skill['category']} | {agents} | "
            f"{meta['yes'] if skill['deletable'] else meta['no']} | {skill['resolved_path']}",
            size=9,
            gap=5,
        )

    out = pdf.output()
    return bytes(out)


def build_docx_report(data: dict[str, Any] | None = None, lang: str = "zh") -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ValueError("Word 导出需要安装 python-docx：pip install python-docx") from exc

    data = data or scan_all()
    meta = _report_meta(lang, data)
    analysis = _build_analysis(lang, data)

    doc = Document()
    title = doc.add_heading(meta["title"], level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    meta_lines = [
        f"{meta['generated']}: {data['scanned_at']}",
        f"{meta['home']}: {data['home']}",
        f"{meta['project']}: {GITHUB_URL}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

    doc.add_heading(meta["summary"], level=1)
    summary_values = {
        "total_skills": data["totals"]["skills"],
        "configured_agents": data["totals"]["agents_configured"],
        "shared_skills": data["totals"]["shared_skills"],
        "disabled": data["totals"]["disabled"],
    }
    for key, value in summary_values.items():
        doc.add_paragraph(f"{meta[key]}: {value}", style="List Bullet")

    doc.add_heading(meta["analysis"], level=1)
    for section in analysis:
        doc.add_heading(section["title"], level=2)
        for paragraph in section["paragraphs"]:
            clean = paragraph.replace("**", "")
            doc.add_paragraph(clean)

    doc.add_heading(meta["agents"], level=1)
    agent_table = doc.add_table(rows=1, cols=4)
    agent_table.style = "Table Grid"
    for idx, header in enumerate(meta["headers_agents"]):
        agent_table.rows[0].cells[idx].text = header
    for agent in data["agents"]:
        row = agent_table.add_row().cells
        row[0].text = agent["label"]
        row[1].text = str(agent["installed"])
        row[2].text = meta["yes"] if agent["configured"] else meta["no"]
        row[3].text = "\n".join(agent["roots"]) or "-"

    doc.add_heading(meta["skills"], level=1)
    skill_table = doc.add_table(rows=1, cols=5)
    skill_table.style = "Table Grid"
    for idx, header in enumerate(meta["headers_skills"]):
        skill_table.rows[0].cells[idx].text = header
    for skill in data["skills"]:
        row = skill_table.add_row().cells
        row[0].text = skill["name"]
        row[1].text = skill["category"]
        row[2].text = ", ".join(skill["agent_labels"]) or "-"
        row[3].text = meta["yes"] if skill["deletable"] else meta["no"]
        row[4].text = skill["resolved_path"]

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_export_bytes(fmt: str, lang: str = "zh", data: dict[str, Any] | None = None) -> tuple[bytes, str, str]:
    fmt = normalize_export_fmt(fmt)
    lang = lang.lower()
    if fmt not in EXPORT_FORMATS:
        raise ValueError(f"format 仅支持 {', '.join(EXPORT_FORMATS)}")
    if lang not in EXPORT_LANGS:
        raise ValueError(f"lang 仅支持 {', '.join(EXPORT_LANGS)}")

    data = data or scan_all()
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    prefix = "skill-report" if lang == "en" else "skill-报告"

    if fmt == "md":
        content = build_markdown_report(data, lang=lang).encode("utf-8")
        return content, "text/markdown; charset=utf-8", f"{prefix}-{stamp}.md"
    if fmt == "docx":
        content = build_docx_report(data, lang=lang)
        return content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{prefix}-{stamp}.docx"
    content = build_pdf_report(data, lang=lang)
    return content, "application/pdf", f"{prefix}-{stamp}.pdf"


def export_report(
    fmt: str = "json",
    output_path: str | None = None,
    lang: str = "en",
) -> dict[str, Any]:
    data = scan_all()
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")

    if fmt == "json":
        content = json.dumps(data, ensure_ascii=False, indent=2)
        suffix = ".json"
        path_content: str | bytes = content
    elif fmt in EXPORT_FORMATS:
        path_content, _, suffix_name = build_export_bytes(fmt, lang, data)
        suffix = Path(suffix_name).suffix
    elif fmt == "markdown":
        content = build_markdown_report(data, lang=lang)
        suffix = ".md"
        path_content = content
    else:
        raise ValueError(f"format 仅支持 json, markdown, {', '.join(EXPORT_FORMATS)}")

    if output_path:
        path = Path(output_path)
    else:
        prefix = "skill-report" if lang == "en" else "skill-报告"
        path = Path.cwd() / f"{prefix}-{timestamp}{suffix}"

    if isinstance(path_content, bytes):
        path.write_bytes(path_content)
    else:
        path.write_text(path_content + ("\n" if fmt in {"markdown", "md"} else ""), encoding="utf-8")

    return {"ok": True, "format": fmt, "lang": lang, "path": str(path.resolve()), "skills": data["totals"]["skills"]}